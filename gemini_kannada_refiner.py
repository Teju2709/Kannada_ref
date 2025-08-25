#!/usr/bin/env python3
"""
Gemini-Enhanced Kannada Text Refinement System
Refines noisy OCR output using Google's Gemini API to produce clean, meaningful Kannada text
"""

import google.generativeai as genai
import os
import json
import time
import re
from typing import List, Dict, Tuple, Optional
from dataclasses import dataclass
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import logging
# ...existing code...
import api  # This will set up the API key
import google.generativeai as genai

# Now you can use genai as needed
# ...existing code...

# Set up logging
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

@dataclass
class RefinementResult:
    original_text: str
    refined_text: str
    confidence_score: float
    corrections_made: List[str]
    processing_time: float

class GeminiKannadaRefiner:
    def __init__(self, api_key: str, model_name: str = "gemini-1.5-pro"):
        """
        Initialize the Gemini-powered Kannada text refiner
        
        Args:
            api_key: Google AI API key
            model_name: Gemini model to use
        """
        genai.configure(api_key=api_key)
        self.model = genai.GenerativeModel(model_name)
        
        # Configure safety settings for text processing
        self.safety_settings = {
            genai.types.HarmCategory.HARM_CATEGORY_HATE_SPEECH: genai.types.HarmBlockThreshold.BLOCK_NONE,
            genai.types.HarmCategory.HARM_CATEGORY_HARASSMENT: genai.types.HarmBlockThreshold.BLOCK_NONE,
            genai.types.HarmCategory.HARM_CATEGORY_SEXUALLY_EXPLICIT: genai.types.HarmBlockThreshold.BLOCK_NONE,
            genai.types.HarmCategory.HARM_CATEGORY_DANGEROUS_CONTENT: genai.types.HarmBlockThreshold.BLOCK_NONE,
        }
        
        # Generation configuration
        self.generation_config = genai.types.GenerationConfig(
            temperature=0.1,  # Low temperature for more consistent corrections
            top_p=0.8,
            top_k=40,
            max_output_tokens=4000,
        )
    
    def create_refinement_prompt(self, noisy_text: str, context: str = None) -> str:
        """
        Create a detailed prompt for Gemini to refine Kannada text
        
        Args:
            noisy_text: OCR output text to be refined
            context: Optional context about the document type
            
        Returns:
            Formatted prompt string
        """
        base_prompt = """
ನೀವು ಕನ್ನಡ ಭಾಷೆಯ ಪರಿಣಿತರು. OCR ತಂತ್ರಜ್ಞಾನದಿಂದ ಬಂದ ಶಬ್ದಯುಕ್ತ ಮತ್ತು ದೋಷಪೂರಿತ ಕನ್ನಡ ಪಠ್ಯವನ್ನು ಸರಿಪಡಿಸಬೇಕು.

**ನಿಮ್ಮ ಕಾರ್ಯ:**
1. ಎಲ್ಲಾ OCR ದೋಷಗಳನ್ನು ಸರಿಪಡಿಸಿ (ತಪ್ಪು ಅಕ್ಷರಗಳು, ಮಿಸ್ಸಿಂಗ್ ಅಕ್ಷರಗಳು)
2. ಅಪೂರ್ಣ ಪದಗಳನ್ನು ಪೂರ್ಣಗೊಳಿಸಿ
3. ವಾಕ್ಯಗಳನ್ನು ವ್ಯಾಕರಣಬದ್ಧವಾಗಿ ಸರಿಪಡಿಸಿ
4. ಅರ್ಥಪೂರ್ಣ ಮತ್ತು ಸುಲಭವಾಗಿ ಓದಬಹುದಾದ ಪಠ್ಯವನ್ನು ರಚಿಸಿ
5. ಮೂಲ ಅರ್ಥವನ್ನು ಕಳೆದುಕೊಳ್ಳಬೇಡಿ

**ಮಾರ್ಗದರ್ಶನಗಳು:**
- ಕೇವಲ ಕನ್ನಡದಲ್ಲಿ ಉತ್ತರಿಸಿ
- ಸಂದರ್ಭವನ್ನು ಆಧರಿಸಿ ಸಮಂಜಸವಾದ ಅನುಮಾನಗಳನ್ನು ಮಾಡಿ
- ಸಾಂಪ್ರದಾಯಿಕ ಕನ್ನಡ ವ್ಯಾಕರಣವನ್ನು ಅನುಸರಿಸಿ
- ಸ್ಪಷ್ಟ ಮತ್ತು ನೈಸರ್ಗಿಕ ಭಾಷೆ ಬಳಸಿ

**ಇನ್‌ಪುಟ್ ಪಠ್ಯ:**
{noisy_text}

**ಸರಿಪಡಿಸಿದ ಪಠ್ಯ:**
"""
        
        if context:
            context_addition = f"\n**ದಾಖಲೆಯ ಸಂದರ್ಭ:** {context}\n"
            base_prompt = base_prompt.replace("**ಇನ್‌ಪುಟ್ ಪಠ್ಯ:**", 
                                            f"**ದಾಖಲೆಯ ಸಂದರ್ಭ:** {context}\n\n**ಇನ್‌ಪುಟ್ ಪಠ್ಯ:**")
        
        return base_prompt.format(noisy_text=noisy_text)
    
    def refine_text_segment(self, text: str, context: str = None, retry_count: int = 3) -> RefinementResult:
        """
        Refine a single text segment using Gemini
        
        Args:
            text: Text segment to refine
            context: Optional context information
            retry_count: Number of retries on failure
            
        Returns:
            RefinementResult object
        """
        start_time = time.time()
        
        for attempt in range(retry_count):
            try:
                # Create prompt
                prompt = self.create_refinement_prompt(text, context)
                
                # Generate refined text
                response = self.model.generate_content(
                    prompt,
                    generation_config=self.generation_config,
                    safety_settings=self.safety_settings
                )
                
                if response.text:
                    refined_text = response.text.strip()
                    processing_time = time.time() - start_time
                    
                    # Calculate simple confidence score based on text improvement
                    confidence_score = self._calculate_confidence(text, refined_text)
                    
                    # Identify corrections made
                    corrections = self._identify_corrections(text, refined_text)
                    
                    return RefinementResult(
                        original_text=text,
                        refined_text=refined_text,
                        confidence_score=confidence_score,
                        corrections_made=corrections,
                        processing_time=processing_time
                    )
                
            except Exception as e:
                logger.error(f"Gemini API error (attempt {attempt + 1}): {e}")
                if attempt < retry_count - 1:
                    time.sleep(2 ** attempt)  # Exponential backoff
                else:
                    # Return original text if all attempts fail
                    return RefinementResult(
                        original_text=text,
                        refined_text=text,
                        confidence_score=0.0,
                        corrections_made=["API Error - No refinement applied"],
                        processing_time=time.time() - start_time
                    )
        
        return RefinementResult(
            original_text=text,
            refined_text=text,
            confidence_score=0.0,
            corrections_made=["Failed to refine"],
            processing_time=time.time() - start_time
        )
    
    def _calculate_confidence(self, original: str, refined: str) -> float:
        """
        Calculate confidence score based on text improvement metrics
        
        Args:
            original: Original noisy text
            refined: Refined text
            
        Returns:
            Confidence score between 0 and 1
        """
        # Simple heuristics for confidence calculation
        score = 0.5  # Base score
        
        # Length improvement (refined should be similar or longer)
        if len(refined) >= len(original):
            score += 0.1
        
        # Kannada character ratio
        kannada_chars_original = len(re.findall(r'[\u0C80-\u0CFF]', original))
        kannada_chars_refined = len(re.findall(r'[\u0C80-\u0CFF]', refined))
        
        if kannada_chars_refined > kannada_chars_original:
            score += 0.2
        
        # Word count improvement
        words_original = len(re.findall(r'[\u0C80-\u0CFF]+', original))
        words_refined = len(re.findall(r'[\u0C80-\u0CFF]+', refined))
        
        if words_refined >= words_original:
            score += 0.1
        
        # Sentence structure (presence of sentence ending marks)
        sentences_refined = len(re.findall(r'[.।!?]', refined))
        if sentences_refined > 0:
            score += 0.1
        
        return min(score, 1.0)
    
    def _identify_corrections(self, original: str, refined: str) -> List[str]:
        """
        Identify types of corrections made
        
        Args:
            original: Original text
            refined: Refined text
            
        Returns:
            List of correction descriptions
        """
        corrections = []
        
        # Character additions
        if len(refined) > len(original):
            corrections.append(f"Added {len(refined) - len(original)} characters")
        
        # Word improvements
        original_words = set(re.findall(r'[\u0C80-\u0CFF]+', original))
        refined_words = set(re.findall(r'[\u0C80-\u0CFF]+', refined))
        
        new_words = refined_words - original_words
        if new_words:
            corrections.append(f"Added/corrected {len(new_words)} words")
        
        # Punctuation improvements
        original_punct = len(re.findall(r'[.।!?,:;]', original))
        refined_punct = len(re.findall(r'[.।!?,:;]', refined))
        
        if refined_punct > original_punct:
            corrections.append("Improved punctuation")
        
        if not corrections:
            corrections.append("Minor refinements")
        
        return corrections
    
    def refine_full_document(self, input_file: str, output_base: str, 
                           chunk_size: int = 1000, context: str = None) -> Dict:
        """
        Refine entire document by processing it in chunks
        
        Args:
            input_file: Path to input text file
            output_base: Base name for output files
            chunk_size: Size of text chunks to process
            context: Document context information
            
        Returns:
            Processing summary dictionary
        """
        logger.info(f"Starting document refinement: {input_file}")
        
        # Read input file
        try:
            with open(input_file, 'r', encoding='utf-8') as f:
                text = f.read()
        except Exception as e:
            logger.error(f"Error reading input file: {e}")
            return {"error": str(e)}
        
        # Split into manageable chunks
        chunks = self._split_into_chunks(text, chunk_size)
        logger.info(f"Split document into {len(chunks)} chunks")
        
        # Process each chunk
        refined_chunks = []
        total_corrections = []
        total_processing_time = 0
        
        for i, chunk in enumerate(chunks):
            logger.info(f"Processing chunk {i + 1}/{len(chunks)}")
            
            result = self.refine_text_segment(chunk, context)
            refined_chunks.append(result.refined_text)
            total_corrections.extend(result.corrections_made)
            total_processing_time += result.processing_time
            
            # Add delay between API calls to avoid rate limiting
            if i < len(chunks) - 1:
                time.sleep(1)
        
        # Combine refined chunks
        refined_document = '\n\n'.join(refined_chunks)
        
        # Save outputs
        self._save_refined_outputs(refined_document, output_base, {
            'original_length': len(text),
            'refined_length': len(refined_document),
            'chunks_processed': len(chunks),
            'total_corrections': len(total_corrections),
            'processing_time': total_processing_time,
            'corrections_summary': list(set(total_corrections))
        })
        
        summary = {
            'input_file': input_file,
            'output_files': [f"{output_base}.txt", f"{output_base}.docx", f"{output_base}_summary.json"],
            'original_length': len(text),
            'refined_length': len(refined_document),
            'chunks_processed': len(chunks),
            'total_corrections': len(total_corrections),
            'processing_time': total_processing_time,
            'improvement_ratio': len(refined_document) / len(text) if len(text) > 0 else 0
        }
        
        logger.info("Document refinement completed successfully!")
        return summary
    
    def _split_into_chunks(self, text: str, chunk_size: int) -> List[str]:
        """
        Split text into manageable chunks for processing
        
        Args:
            text: Input text
            chunk_size: Target size for each chunk
            
        Returns:
            List of text chunks
        """
        # Split by paragraphs first
        paragraphs = re.split(r'\n\s*\n', text)
        chunks = []
        current_chunk = ""
        
        for para in paragraphs:
            if len(current_chunk) + len(para) <= chunk_size:
                current_chunk += para + "\n\n"
            else:
                if current_chunk:
                    chunks.append(current_chunk.strip())
                current_chunk = para + "\n\n"
        
        # Add the last chunk
        if current_chunk:
            chunks.append(current_chunk.strip())
        
        return chunks
    
    def _save_refined_outputs(self, refined_text: str, output_base: str, metadata: Dict):
        """
        Save refined text in multiple formats with metadata
        
        Args:
            refined_text: Refined text content
            output_base: Base filename for outputs
            metadata: Processing metadata
        """
        # Save as plain text
        with open(f"{output_base}.txt", 'w', encoding='utf-8') as f:
            f.write(refined_text)
        
        # Save as Word document
        try:
            doc = Document()
            
            # Add title
            title = doc.add_heading('ಸರಿಪಡಿಸಿದ ಕನ್ನಡ ದಾಖಲೆ', 0)
            title.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # Add metadata
            doc.add_paragraph(f"ಸರಿಪಡಿಸಿದ ದಿನಾಂಕ: {time.strftime('%Y-%m-%d %H:%M')}")
            doc.add_paragraph(f"ಮೂಲ ಪಠ್ಯದ ಉದ್ದ: {metadata['original_length']} ಅಕ್ಷರಗಳು")
            doc.add_paragraph(f"ಸರಿಪಡಿಸಿದ ಪಠ್ಯದ ಉದ್ದ: {metadata['refined_length']} ಅಕ್ಷರಗಳು")
            doc.add_paragraph(f"ಮಾಡಿದ ಸುಧಾರಣೆಗಳು: {metadata['total_corrections']}")
            doc.add_paragraph()
            
            # Add refined content
            for paragraph in refined_text.split('\n\n'):
                if paragraph.strip():
                    doc.add_paragraph(paragraph.strip())
            
            doc.save(f"{output_base}.docx")
            logger.info(f"Word document saved: {output_base}.docx")
            
        except Exception as e:
            logger.error(f"Error saving Word document: {e}")
        
        # Save processing summary
        summary = {
            **metadata,
            'processing_timestamp': time.strftime('%Y-%m-%d %H:%M:%S'),
            'gemini_model_used': 'gemini-1.5-pro',
            'api_version': 'google-generativeai'
        }
        
        with open(f"{output_base}_summary.json", 'w', encoding='utf-8') as f:
            json.dump(summary, f, indent=2, ensure_ascii=False)
        
        logger.info(f"Processing summary saved: {output_base}_summary.json")

def main():
    """
    Example usage of the Gemini Kannada Refiner
    """
    import argparse
    
    parser = argparse.ArgumentParser(description='Refine Kannada OCR output using Gemini AI')
    parser.add_argument('input_file', help='Input text file from OCR post-processing')
    parser.add_argument('output_base', help='Base name for output files')
    parser.add_argument('--api-key', required=True, help='Google AI API key')
    parser.add_argument('--context', help='Document context (e.g., "historical literature", "newspaper")')
    parser.add_argument('--chunk-size', type=int, default=1000, help='Text chunk size for processing')
    parser.add_argument('--model', default='gemini-1.5-pro', help='Gemini model to use')
    
    args = parser.parse_args()
    
    # Initialize refiner
    refiner = GeminiKannadaRefiner(args.api_key, args.model)
    
    # Process document
    result = refiner.refine_full_document(
        args.input_file, 
        args.output_base, 
        args.chunk_size, 
        args.context
    )
    
    # Print summary
    if 'error' not in result:
        print(f"\n{'='*50}")
        print("REFINEMENT COMPLETED SUCCESSFULLY")
        print(f"{'='*50}")
        print(f"Input: {result['input_file']}")
        print(f"Original length: {result['original_length']} characters")
        print(f"Refined length: {result['refined_length']} characters")
        print(f"Improvement ratio: {result['improvement_ratio']:.2f}x")
        print(f"Chunks processed: {result['chunks_processed']}")
        print(f"Total corrections: {result['total_corrections']}")
        print(f"Processing time: {result['processing_time']:.2f} seconds")
        print(f"\nOutput files:")
        for file in result['output_files']:
            print(f"  - {file}")
    else:
        print(f"Error: {result['error']}")

if __name__ == "__main__":
    # Example without command line arguments
    if len(os.sys.argv) == 1:
        print("Example usage:")
        print("python gemini_kannada_refiner.py extracted_text.txt refined_output --api-key YOUR_API_KEY")
        print("\nTo get a Google AI API key:")
        print("1. Go to https://makersuite.google.com/app/apikey")
        print("2. Create a new API key")
        print("3. Use it with --api-key parameter")
    else:
        main()