#!/usr/bin/env python3
"""
Kannada OCR Post-Processing System
Comprehensive text cleanup, spell checking, and formatting
"""

#!/usr/bin/env python3
"""
Kannada OCR Post-Processing System
Comprehensive text cleanup, spell checking, and formatting
"""

import re
import json
import os
from typing import List, Dict, Tuple
from dataclasses import dataclass
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
import requests
import time
import openai
# ...existing code...
import api  # This will set up the API key
import google.generativeai as genai

# Now you can use genai as needed
# ...existing code...

@dataclass
class TextSegment:
    text: str
    segment_type: str  # 'paragraph', 'heading', 'poem', 'list'
    confidence: float = 0.0
    original_text: str = ""


class KannadaPostProcessor:
    def __init__(self, use_openai=False, openai_api_key=None):
        self.use_openai = use_openai
        if openai_api_key:
            openai.api_key = openai_api_key

        # Moved the below block from inside correct_with_gpt to __init__
        self.char_corrections = {
            # Similar looking characters
            'ಙ': 'ಞ',  # Common confusion
            'ಛ': 'ಚ',  # Similar shapes
            'ಠ': 'ಟ',  # OCR confusion
            'ಧ': 'ದ',  # Similar strokes
            'ಭ': 'ಬ',  # Similar shapes
            'ೞ': 'ಳ',  # Archaic vs modern
            'ೠ': 'ೃ',  # Vowel marks
            '಻': '಼',  # Diacritical marks
            # Common punctuation fixes
            '।': '.',  # Devanagari danda to period
            '॥': '.',  # Double danda
        }

        # Word-level corrections (common OCR mistakes)
        self.word_corrections = {
            'ಕನ್ನಡ': ['ಕನ್ನಡ್', 'ಕನ್ನಡಾ', 'ಕನ್ನಡ಼'],
            'ಬರೆಯುವ': ['ಬರೆಯುವ್', 'ಬರೆಯ್ುವ'],
            'ಮಾಡುವ': ['ಮಾಡ್ುವ', 'ಮಾಡುವ್'],
            'ಹೇಳುವ': ['ಹೇಳ್ುವ', 'ಹೇಳುವ್'],
            'ಮನುಷ್ಯ': ['ಮನುಷ್ಯ್', 'ಮನುಷ್ಯಾ'],
        }

        # Regex patterns for cleanup
        self.cleanup_patterns = [
            (r'\s+', ' '),
            (r'([ಕ-ಹ])\s+([ಾ-ೌ])', r'\1\2'),
            (r'[^\u0C80-\u0CFF\u0020-\u007E\n]', ''),
            (r'\s*([.!?,:;])\s*', r'\1 '),
            (r'\s+$', ''),
            (r'\n\s*\n\s*\n+', '\n\n'),
        ]

        self.stopwords = self._load_kannada_stopwords()
        self.common_words = self._load_common_kannada_words()

    # ... rest of the class remains unchanged ...

# No changes needed in the rest of the class or in main()
# Ensure the char_corrections dictionary is initialized during __init__ and not inside any method

    def _load_kannada_stopwords(self) -> set:
        """Load Kannada stopwords"""
        # Common Kannada stopwords
        stopwords = {
            'ಮತ್ತು', 'ಅಥವಾ', 'ಆದರೆ', 'ಅಲ್ಲದೆ', 'ಇಲ್ಲಿ', 'ಅಲ್ಲಿ', 
            'ಎಲ್ಲಿ', 'ಯಾವಾಗ', 'ಏಕೆ', 'ಹೇಗೆ', 'ಯಾರು', 'ಏನು', 'ಯಾವ',
            'ಇದು', 'ಅದು', 'ಆತ', 'ಇವನು', 'ಅವನು', 'ಇವಳು', 'ಅವಳು',
            'ನಾನು', 'ನೀನು', 'ಅವನು', 'ಅವಳು', 'ನಾವು', 'ನೀವು', 'ಅವರು'
        }
        return stopwords
    
    def _load_common_kannada_words(self) -> set:
        """Load common Kannada words for spell checking"""
        # This would ideally load from a dictionary file
        # For now, using a small set of common words
        common_words = {
            'ಕನ್ನಡ', 'ಭಾಷೆ', 'ಪುಸ್ತಕ', 'ಮನೆ', 'ಊರು', 'ದೇಶ', 'ಜನ',
            'ಜೀವನ', 'ಪ್ರೇಮ', 'ಶಿಕ್ಷಣ', 'ಕಲೆ', 'ಸಂಸ್ಕೃತಿ', 'ಇತಿಹಾಸ',
            'ಪ್ರಕೃತಿ', 'ಸಮಾಜ', 'ರಾಜಕೀಯ', 'ವಿಜ್ಞಾನ', 'ತಂತ್ರಜ್ಞಾನ'
        }
        return common_words
    
    def clean_raw_ocr(self, text: str) -> str:
        """
        Clean raw OCR output with basic corrections
        
        Args:
            text: Raw OCR text
            
        Returns:
            Cleaned text
        """
        cleaned = text
        
        # Apply character-level corrections
        for wrong, correct in self.char_corrections.items():
            cleaned = cleaned.replace(wrong, correct)
        
        # Apply regex cleanup patterns
        for pattern, replacement in self.cleanup_patterns:
            cleaned = re.sub(pattern, replacement, cleaned)
        
        return cleaned.strip()
    
    def fix_broken_words(self, text: str) -> str:
        """
        Fix words broken across lines
        
        Args:
            text: Input text
            
        Returns:
            Text with fixed word breaks
        """
        lines = text.split('\n')
        fixed_lines = []
        
        for i, line in enumerate(lines):
            line = line.strip()
            if not line:
                fixed_lines.append(line)
                continue
            
            # Check if line ends with a consonant (likely broken word)
            if i < len(lines) - 1 and re.search(r'[ಕ-ಹ]$', line):
                next_line = lines[i + 1].strip()
                # If next line starts with vowel mark, merge
                if re.search(r'^[ಾ-ೌ]', next_line):
                    line += next_line
                    lines[i + 1] = ""  # Mark for removal
            
            fixed_lines.append(line)
        
        # Remove empty lines created by merging
        return '\n'.join([line for line in fixed_lines if line.strip()])
    
    def spell_check_and_correct(self, text: str) -> Tuple[str, List[Dict]]:
        """
        Spell check and correct using word patterns
        
        Args:
            text: Input text
            
        Returns:
            Tuple of (corrected_text, list_of_corrections)
        """
        corrections = []
        words = re.findall(r'[\u0C80-\u0CFF]+', text)  # Extract Kannada words
        corrected_text = text
        
        for word in words:
            # Check against known corrections
            for correct_word, variations in self.word_corrections.items():
                if word in variations:
                    corrected_text = corrected_text.replace(word, correct_word)
                    corrections.append({
                        'original': word,
                        'corrected': correct_word,
                        'confidence': 0.8
                    })
        
        return corrected_text, corrections
    
    def detect_text_structure(self, text: str) -> List[TextSegment]:
        """
        Detect and classify text segments (headings, paragraphs, poems, etc.)
        
        Args:
            text: Input text
            
        Returns:
            List of TextSegment objects
        """
        segments = []
        paragraphs = re.split(r'\n\s*\n', text)
        
        for para in paragraphs:
            para = para.strip()
            if not para:
                continue
            
            segment_type = self._classify_text_segment(para)
            segments.append(TextSegment(
                text=para,
                segment_type=segment_type,
                confidence=0.7,
                original_text=para
            ))
        
        return segments
    
    def _classify_text_segment(self, text: str) -> str:
        """
        Classify a text segment into type
        
        Args:
            text: Text segment
            
        Returns:
            Segment type
        """
        lines = text.split('\n')
        
        # Single line, short text - likely heading
        if len(lines) == 1 and len(text) < 100:
            # Check for heading indicators
            if re.search(r'^[ಅ-ಹ].{10,50}$', text):
                return 'heading'
        
        # Multiple short lines - likely poem
        if len(lines) > 2 and all(len(line.strip()) < 50 for line in lines):
            return 'poem'
        
        # Lists (numbered or bulleted)
        if re.search(r'^\s*[೧-೯0-9]\.', text) or re.search(r'^\s*[•\-\*]', text):
            return 'list'
        
        # Default to paragraph
        return 'paragraph'
    
    def ai_assisted_correction(self, text: str, use_openai: bool = False) -> str:
        """
        Use AI for context-based correction (placeholder for API integration)
        
        Args:
            text: Input text
            use_openai: Whether to use OpenAI API
            
        Returns:
            AI-corrected text
        """
        if use_openai:
            # This would integrate with OpenAI API
            # For now, return basic corrections
            return self._basic_ai_corrections(text)
        else:
            return self._basic_ai_corrections(text)
    
    def _basic_ai_corrections(self, text: str) -> str:
        """
        Basic AI-like corrections using patterns
        
        Args:
            text: Input text
            
        Returns:
            Corrected text
        """
        # Context-based corrections
        corrections = {
            # Fix common contextual errors
            r'ಕನ್ನಡ\s+ಭಾಶೆ': 'ಕನ್ನಡ ಭಾಷೆ',
            r'ಪುಸ್ತಕ\s+ವನ್ನು': 'ಪುಸ್ತಕವನ್ನು',
            r'ಮನೆ\s+ಗೆ': 'ಮನೆಗೆ',
            r'ಊರು\s+ಇಂದ': 'ಊರಿಂದ',
        }
        
        corrected = text
        for pattern, replacement in corrections.items():
            corrected = re.sub(pattern, replacement, corrected)
        
        return corrected
    
    def format_as_document(self, segments: List[TextSegment]) -> Document:
        """
        Create a formatted Word document from text segments
        
        Args:
            segments: List of text segments
            
        Returns:
            python-docx Document object
        """
        doc = Document()
        
        # Add title
        title = doc.add_heading('ಪರಿವರ್ತಿತ ಕನ್ನಡ ದಾಖಲೆ', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Add metadata
        doc.add_paragraph(f"ಪರಿವರ್ತನೆ ದಿನಾಂಕ: {time.strftime('%Y-%m-%d %H:%M')}")
        doc.add_paragraph()
        
        for segment in segments:
            if segment.segment_type == 'heading':
                heading = doc.add_heading(segment.text, 1)
                heading.alignment = WD_ALIGN_PARAGRAPH.LEFT
            
            elif segment.segment_type == 'poem':
                # Add poem with special formatting
                poem_para = doc.add_paragraph()
                poem_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                poem_para.add_run(segment.text).italic = True
                
            elif segment.segment_type == 'list':
                # Format as bullet list
                for line in segment.text.split('\n'):
                    if line.strip():
                        doc.add_paragraph(line.strip(), style='List Bullet')
            
            else:  # paragraph
                doc.add_paragraph(segment.text)
        
        return doc
    
    def save_formatted_output(self, segments: List[TextSegment], base_filename: str):
        """
        Save processed text in multiple formats
        
        Args:
            segments: Processed text segments
            base_filename: Base filename without extension
        """
        # Save as plain text
        with open(f"{base_filename}.txt", 'w', encoding='utf-8') as f:
            for segment in segments:
                if segment.segment_type == 'heading':
                    f.write(f"\n# {segment.text}\n\n")
                elif segment.segment_type == 'poem':
                    f.write(f"\n*{segment.text}*\n\n")
                else:
                    f.write(f"{segment.text}\n\n")
        
        # Save as Markdown
        with open(f"{base_filename}.md", 'w', encoding='utf-8') as f:
            for segment in segments:
                if segment.segment_type == 'heading':
                    f.write(f"\n# {segment.text}\n\n")
                elif segment.segment_type == 'poem':
                    f.write(f"\n*{segment.text}*\n\n")
                elif segment.segment_type == 'list':
                    for line in segment.text.split('\n'):
                        if line.strip():
                            f.write(f"- {line.strip()}\n")
                    f.write("\n")
                else:
                    f.write(f"{segment.text}\n\n")
        
        # Save as Word document
        try:
            doc = self.format_as_document(segments)
            doc.save(f"{base_filename}.docx")
            print(f"✓ Saved formatted document: {base_filename}.docx")
        except Exception as e:
            print(f"Warning: Could not save Word document: {e}")
        
        # Save processing report
        report = {
            'total_segments': len(segments),
            'segment_types': {seg_type: len([s for s in segments if s.segment_type == seg_type]) 
                            for seg_type in ['paragraph', 'heading', 'poem', 'list']},
            'processing_time': time.strftime('%Y-%m-%d %H:%M:%S')
        }
        
        with open(f"{base_filename}_report.json", 'w', encoding='utf-8') as f:
            json.dump(report, f, indent=2, ensure_ascii=False)
    
    def process_full_pipeline(self, raw_ocr_text: str, output_base: str) -> Dict:
        """
        Complete post-processing pipeline
        
        Args:
            raw_ocr_text: Raw OCR output
            output_base: Base filename for outputs
            
        Returns:
            Processing results dictionary
        """
        print("Starting Kannada OCR post-processing...")
        
        # Step 1: Basic cleanup
        print("Step 1: Basic text cleanup...")
        cleaned_text = self.clean_raw_ocr(raw_ocr_text)
        
        # Step 2: Fix broken words
        print("Step 2: Fixing broken words...")
        fixed_text = self.fix_broken_words(cleaned_text)
        
        # Step 3: Spell checking
        print("Step 3: Spell checking and corrections...")
        corrected_text, corrections = self.spell_check_and_correct(fixed_text)
        
        # Step 4: AI-assisted corrections
        print("Step 4: AI-assisted corrections...")
        ai_corrected_text = self.ai_assisted_correction(corrected_text)
        
        # Step 5: Structure detection
        print("Step 5: Detecting text structure...")
        segments = self.detect_text_structure(ai_corrected_text)
        
        # Step 6: Save formatted outputs
        print("Step 6: Saving formatted outputs...")
        self.save_formatted_output(segments, output_base)
        
        # Return results
        results = {
            'original_length': len(raw_ocr_text),
            'processed_length': len(ai_corrected_text),
            'corrections_made': len(corrections),
            'segments_detected': len(segments),
            'segment_breakdown': {
                seg_type: len([s for s in segments if s.segment_type == seg_type])
                for seg_type in ['paragraph', 'heading', 'poem', 'list']
            }
        }
        
        print(f"✓ Post-processing complete!")
        print(f"  - Made {results['corrections_made']} corrections")
        print(f"  - Detected {results['segments_detected']} text segments")
        print(f"  - Saved outputs as: {output_base}.[txt|md|docx]")
        
        return results

def main():
    
    """Example usage of the post-processor"""
    # Example raw OCR text (replace with your actual OCR output)
    sample_text = """
    ಕನ್ನಡ ಸಾಹಿತ್ಯ
    
    ಕನ್ನಡ ಭಾಷೆಯ ಸಾಹಿತ್ಯವು ಬಹಳ ಸಮೃದ್ಧವಾಗಿದೆ. ಇದು ಸಾವಿರಾರು 
    ವರ್ಷಗಳ ಇತಿಹಾಸವನ್ನು ಹೊಂದಿದೆ.
    
    ಪ್ರಮುಖ ಕವಿಗಳು:
    ೧. ಪಂಪ
    ೨. ರನ್ನ  
    ೩. ಪೊನ್ನ
    
    ಕವನ ಉದಾಹರಣೆ
    ಜಗವೆಲ್ಲಾ ಇಂದು ಹಬ್ಬದ ದಿನ
    ಮನವೆಲ್ಲಾ ಇಂದು ಉತ್ಸಾಹದಿಂ
    """
    
    # Initialize post-processor
    processor = KannadaPostProcessor()
    
    # Run full pipeline
    results = processor.process_full_pipeline(sample_text, "processed_kannada_text")
    
    print(f"\nProcessing Results: {results}")

if __name__ == "__main__":
    main()
    # Individual processing steps



